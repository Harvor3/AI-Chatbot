import re
import json
import streamlit as st
from typing import Dict, Any, Optional, List
from langchain_core.language_models import BaseLanguageModel
from langchain.prompts import ChatPromptTemplate
from .base_agent import BaseAgent
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import io
import tempfile
import os


class FormGenerationAgent(BaseAgent):
    def __init__(self, llm: BaseLanguageModel):
        super().__init__(
            name="Form Generation Agent",
            description="Creates dynamic forms and handles form-related requests"
        )
        self.llm = llm
        
        self.prompt = ChatPromptTemplate.from_messages([
            ("system", """You are a specialized form generation assistant. Your role is to:
            1. Create dynamic forms based on user requirements
            2. Generate form schemas in JSON format for actual form building
            3. Suggest appropriate field types and layouts
            4. Handle form validation and error handling
            5. Provide form templates that can be rendered as actual forms
            
            IMPORTANT: When creating forms, respond with a valid JSON schema in this exact format:
            {{
                "form_title": "Form Title",
                "form_description": "Brief description",
                "fields": [
                    {{
                        "id": "field_id",
                        "type": "text|email|number|textarea|select|radio|checkbox|date|file",
                        "label": "Field Label",
                        "placeholder": "Placeholder text",
                        "required": true|false,
                        "options": ["option1", "option2"] (only for select/radio),
                        "validation": "validation rules"
                    }}
                ]
            }}
            
            Always create user-friendly forms with proper validation and clear labels."""),
            ("human", "{message}")
        ])
    
    def process(self, message: str, context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        try:
            chain = self.prompt | self.llm
            
            response = chain.invoke({"message": message})
            response_content = response.content if hasattr(response, 'content') else str(response)
            
            # Try to extract JSON schema from response
            form_schema = self._extract_form_schema(response_content)
            
            result = {
                "response": response_content,
                "agent": self.name,
                "confidence": self.can_handle(message, context),
                "metadata": {
                    "type": "form_generation",
                    "generates_ui_components": True,
                    "form_schema": form_schema
                }
            }
            
            # If we have a valid schema, offer to build the form
            if form_schema:
                result["metadata"]["can_build_form"] = True
                result["response"] = f"✅ **Form Schema Generated Successfully!**\n\n{response_content}\n\n🎯 **Ready to build interactive form with preview and export options!**"
            
            return result
            
        except Exception as e:
            return {
                "response": f"I encountered an error generating your form: {str(e)}",
                "agent": self.name,
                "confidence": 0.0,
                "error": str(e)
            }
    
    def can_handle(self, message: str, context: Optional[Dict[str, Any]] = None) -> float:
        message_lower = message.lower()
        
        high_confidence_keywords = [
            "form", "create form", "generate form", "form builder", "input field",
            "form validation", "form schema", "contact form", "registration form",
            "survey form", "feedback form"
        ]
        
        medium_confidence_keywords = [
            "field", "input", "validation", "schema", "template", "layout",
            "checkbox", "radio button", "dropdown", "text field", "submit"
        ]
        
        for keyword in high_confidence_keywords:
            if keyword in message_lower:
                return 0.9
        
        for keyword in medium_confidence_keywords:
            if keyword in message_lower:
                return 0.6
        
        if re.search(r'\b(html|css|javascript|react|vue|angular)\b', message_lower):
            return 0.7
        
        return 0.1
    
    def _extract_form_schema(self, response_content: str) -> Optional[Dict]:
        """Extract JSON form schema from LLM response"""
        try:
            # Look for JSON content in the response
            import re
            json_match = re.search(r'\{.*\}', response_content, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                schema = json.loads(json_str)
                
                # Validate required fields
                if "form_title" in schema and "fields" in schema:
                    return schema
            return None
        except:
            return None
    
    def build_streamlit_form(self, form_schema: Dict) -> None:
        """Build interactive Streamlit form from schema"""
        st.subheader(f"📝 {form_schema.get('form_title', 'Generated Form')}")
        
        if form_schema.get('form_description'):
            st.write(form_schema['form_description'])
        
        # Initialize form data in session state if not exists
        form_key = f"form_data_{form_schema.get('form_title', 'default').replace(' ', '_')}"
        if form_key not in st.session_state:
            st.session_state[form_key] = {}
        
        # Create unique form key and store it in session state
        form_ui_key = f"form_ui_key_{form_key}"
        if form_ui_key not in st.session_state:
            import uuid
            st.session_state[form_ui_key] = f"form_{uuid.uuid4().hex[:8]}"
        
        form_data = {}
        
        with st.form(key=st.session_state[form_ui_key]):
            for field in form_schema.get('fields', []):
                field_id = field.get('id', 'field')
                field_type = field.get('type', 'text')
                label = field.get('label', 'Field')
                placeholder = field.get('placeholder', '')
                required = field.get('required', False)
                options = field.get('options', [])
                
                # Get stored value or default
                stored_value = st.session_state[form_key].get(field_id, None)
                
                # Add required indicator
                if required:
                    label += " *"
                
                # Render different field types with stored values
                if field_type == 'text':
                    form_data[field_id] = st.text_input(label, value=stored_value or "", placeholder=placeholder)
                elif field_type == 'email':
                    form_data[field_id] = st.text_input(label, value=stored_value or "", placeholder=placeholder or "example@email.com")
                elif field_type == 'number':
                    form_data[field_id] = st.number_input(label, value=stored_value or 0)
                elif field_type == 'textarea':
                    form_data[field_id] = st.text_area(label, value=stored_value or "", placeholder=placeholder)
                elif field_type == 'select':
                    default_index = 0
                    if stored_value and stored_value in options:
                        default_index = options.index(stored_value)
                    form_data[field_id] = st.selectbox(label, options, index=default_index)
                elif field_type == 'radio':
                    default_index = 0
                    if stored_value and stored_value in options:
                        default_index = options.index(stored_value)
                    form_data[field_id] = st.radio(label, options, index=default_index)
                elif field_type == 'checkbox':
                    if options:
                        form_data[field_id] = st.multiselect(label, options, default=stored_value or [])
                    else:
                        form_data[field_id] = st.checkbox(label, value=stored_value or False)
                elif field_type == 'date':
                    form_data[field_id] = st.date_input(label, value=stored_value)
                elif field_type == 'file':
                    form_data[field_id] = st.file_uploader(label)
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                submitted = st.form_submit_button("📋 Submit Form")
            with col2:
                export_pdf = st.form_submit_button("📄 Export PDF")
            with col3:
                export_docx = st.form_submit_button("📝 Export DOCX")
            with col4:
                clear_form = st.form_submit_button("🧹 Clear Data")
            
            # Handle form actions
            if submitted or export_pdf or export_docx or clear_form:
                # Store form data
                st.session_state[form_key] = form_data
                
                if submitted:
                    st.success("✅ Form submitted successfully!")
                    st.json(form_data)
                
                elif export_pdf:
                    st.session_state[f"{form_key}_export_pdf"] = True
                    st.success("✅ Preparing PDF export...")
                
                elif export_docx:
                    st.session_state[f"{form_key}_export_docx"] = True
                    st.success("✅ Preparing DOCX export...")
                
                elif clear_form:
                    st.session_state[form_key] = {}
                    st.success("🧹 Form data cleared!")
                    st.rerun()
        
        # Handle export actions outside the form
        if st.session_state.get(f"{form_key}_export_pdf"):
            try:
                pdf_buffer = self.export_to_pdf(form_schema, st.session_state[form_key])
                st.download_button(
                    label="⬇️ Download PDF Form",
                    data=pdf_buffer,
                    file_name=f"{form_schema.get('form_title', 'form').replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    key=f"pdf_download_{form_key}"
                )
                # Clear the export flag
                del st.session_state[f"{form_key}_export_pdf"]
            except Exception as e:
                st.error(f"❌ Error generating PDF: {str(e)}")
                del st.session_state[f"{form_key}_export_pdf"]
        
        if st.session_state.get(f"{form_key}_export_docx"):
            try:
                docx_buffer = self.export_to_docx(form_schema, st.session_state[form_key])
                st.download_button(
                    label="⬇️ Download DOCX Form",
                    data=docx_buffer,
                    file_name=f"{form_schema.get('form_title', 'form').replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"docx_download_{form_key}"
                )
                # Clear the export flag
                del st.session_state[f"{form_key}_export_docx"]
            except Exception as e:
                st.error(f"❌ Error generating DOCX: {str(e)}")
                del st.session_state[f"{form_key}_export_docx"]
    
    def export_to_pdf(self, form_schema: Dict, form_data: Dict) -> bytes:
        """Export form to PDF"""
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        
        # Title
        p.setFont("Helvetica-Bold", 16)
        p.drawString(50, 750, form_schema.get('form_title', 'Generated Form'))
        
        # Description
        if form_schema.get('form_description'):
            p.setFont("Helvetica", 12)
            p.drawString(50, 720, form_schema['form_description'])
        
        # Fields
        y_position = 680
        p.setFont("Helvetica", 12)
        
        for field in form_schema.get('fields', []):
            field_id = field.get('id', 'field')
            label = field.get('label', 'Field')
            value = form_data.get(field_id, '')
            
            # Field label
            p.setFont("Helvetica-Bold", 11)
            p.drawString(50, y_position, f"{label}:")
            
            # Field value/space
            p.setFont("Helvetica", 10)
            if value:
                p.drawString(50, y_position - 20, str(value))
            else:
                # Draw line for empty field
                p.line(50, y_position - 15, 500, y_position - 15)
            
            y_position -= 50
            
            # Start new page if needed
            if y_position < 100:
                p.showPage()
                y_position = 750
        
        p.save()
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_docx(self, form_schema: Dict, form_data: Dict) -> bytes:
        """Export form to DOCX"""
        doc = Document()
        
        # Title
        title = doc.add_heading(form_schema.get('form_title', 'Generated Form'), 0)
        
        # Description
        if form_schema.get('form_description'):
            doc.add_paragraph(form_schema['form_description'])
        
        # Fields
        for field in form_schema.get('fields', []):
            field_id = field.get('id', 'field')
            label = field.get('label', 'Field')
            value = form_data.get(field_id, '')
            
            # Field label
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            
            # Field value or space
            if value:
                p.add_run(str(value))
            else:
                p.add_run("_" * 50)  # Placeholder line
            
            doc.add_paragraph()  # Space between fields
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue() 